from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph

from src.config import OUTPUT_DIR, TEMPLATES_DIR
from src.database import Database
from src.logger import get_logger
from src.models import WorshipBulletin


logger = get_logger(__name__)


class BulletinTemplateError(RuntimeError):
    """Erro amigavel relacionado aos modelos DOCX de boletim."""


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", value or "sem_data").strip("_")
    return cleaned or "sem_data"


def list_bulletin_templates() -> list[Path]:
    """Lista todos os modelos .docx disponiveis na pasta templates/."""
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    return sorted(
        path
        for path in TEMPLATES_DIR.glob("*.docx")
        if path.is_file() and not path.name.startswith("~$")
    )


def get_template_by_name(template_name: str) -> Path:
    """Busca um modelo pelo nome do arquivo ou caminho informado."""
    clean_name = (template_name or "").strip()
    if not clean_name:
        raise BulletinTemplateError("Nenhum modelo de boletim foi selecionado.")

    direct_path = Path(clean_name)
    candidates = list_bulletin_templates()

    if direct_path.exists() and direct_path.suffix.lower() == ".docx":
        return direct_path

    for template in candidates:
        if template.name == clean_name:
            return template

    for template in candidates:
        if template.stem == clean_name:
            return template

    raise BulletinTemplateError(f'Modelo de boletim nao encontrado: "{clean_name}".')


def resolve_template(template_name: str | None = None) -> Path:
    templates = list_bulletin_templates()
    if template_name:
        return get_template_by_name(template_name)
    if not templates:
        raise BulletinTemplateError(
            "Nenhum modelo de boletim encontrado. Coloque arquivos .docx na pasta templates/."
        )
    if len(templates) > 1:
        raise BulletinTemplateError(
            "Ha mais de um modelo de boletim disponivel. Escolha o tipo de culto antes de gerar o boletim."
        )
    return templates[0]


def build_placeholder_mapping(bulletin: WorshipBulletin) -> dict[str, str]:
    values = {
        "DATA": bulletin.date_text,
        "DIRIGENTE": bulletin.dirigente,
        "PREGADOR": bulletin.pregador,
        "PRELUDIO_MUSICA": bulletin.preludio_musica,
        "PRELUDIO_CANTOR": bulletin.preludio_cantor,
        "PRELUDIO_TOM": bulletin.preludio_tom,
        "MUSICA1": bulletin.musica1,
        "CANTOR1": bulletin.cantor1,
        "TOM1": bulletin.tom1,
        "REF1": bulletin.ref1,
        "TEXTO1": bulletin.texto1,
        "MUSICA2": bulletin.musica2,
        "CANTOR2": bulletin.cantor2,
        "TOM2": bulletin.tom2,
        "REF2": bulletin.ref2,
        "TEXTO2": bulletin.texto2,
        "MUSICA3": bulletin.musica3,
        "CANTOR3": bulletin.cantor3,
        "TOM3": bulletin.tom3,
        "REF3": bulletin.ref3,
        "TEXTO3": bulletin.texto3,
        "ORACAO_LOUVOR": bulletin.oracao_louvor,
        "REF_LOUVOR": bulletin.ref_louvor,
        "TEXTO_LOUVOR": bulletin.texto_louvor,
        "OFERTAS_REF": bulletin.ofertas_ref,
        "OFERTAS_TEXTO": bulletin.ofertas_texto,
        "OFERTAS_ORACAO": bulletin.ofertas_oracao,
        "MUSICA4": bulletin.musica4,
        "CANTOR4": bulletin.cantor4,
        "TOM4": bulletin.tom4,
        "MUSICA5": bulletin.musica5,
        "CANTOR5": bulletin.cantor5,
        "TOM5": bulletin.tom5,
        "ORACAO_INTERCESSAO": bulletin.oracao_intercessao,
        "MUSICA_PAO": bulletin.musica_pao,
        "CANTOR_PAO": bulletin.cantor_pao,
        "TOM_PAO": bulletin.tom_pao,
        "MUSICA_VINHO": bulletin.musica_vinho,
        "CANTOR_VINHO": bulletin.cantor_vinho,
        "TOM_VINHO": bulletin.tom_vinho,
        "MUSICA_EXTRA": bulletin.musica_extra,
        "CANTOR_EXTRA": bulletin.cantor_extra,
        "TOM_EXTRA": bulletin.tom_extra,
        "MUSICA_FINAL": bulletin.musica_final,
        "CANTOR_FINAL": bulletin.cantor_final,
        "TOM_FINAL": bulletin.tom_final,
    }
    return {f"{{{{{key}}}}}": str(value or "") for key, value in values.items()}


def _paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def _replace_in_paragraph(paragraph: Paragraph, mapping: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    original_text = _paragraph_text(paragraph)
    new_text = original_text
    for placeholder, value in mapping.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, value)
    if new_text == original_text:
        return

    first_run = paragraph.runs[0]
    for run in paragraph.runs:
        run.text = ""
    first_run.text = new_text


def _iter_table_paragraphs(container) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
                paragraphs.extend(_iter_table_paragraphs(cell))
    return paragraphs


def replace_placeholders_in_docx(document: DocumentObject, mapping: dict[str, str]) -> None:
    """Substitui placeholders em corpo, tabelas, cabecalhos e rodapes."""
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, mapping)
    for paragraph in _iter_table_paragraphs(document):
        _replace_in_paragraph(paragraph, mapping)

    for section in document.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                _replace_in_paragraph(paragraph, mapping)
            for paragraph in _iter_table_paragraphs(header_footer):
                _replace_in_paragraph(paragraph, mapping)


def find_unreplaced_placeholders(document: DocumentObject) -> list[str]:
    pattern = re.compile(r"\{\{[A-Z0-9_]+\}\}")
    found: set[str] = set()

    def collect_from_paragraphs(paragraphs: list[Paragraph]) -> None:
        for paragraph in paragraphs:
            found.update(pattern.findall(_paragraph_text(paragraph)))

    collect_from_paragraphs(document.paragraphs)
    collect_from_paragraphs(_iter_table_paragraphs(document))
    for section in document.sections:
        collect_from_paragraphs(section.header.paragraphs)
        collect_from_paragraphs(_iter_table_paragraphs(section.header))
        collect_from_paragraphs(section.footer.paragraphs)
        collect_from_paragraphs(_iter_table_paragraphs(section.footer))
    return sorted(found)


def generate_bulletin_docx(
    bulletin_id: int,
    output_path: str | Path | None = None,
    template_name: str | None = None,
) -> Path:
    database = Database()
    bulletin = database.get_bulletin(bulletin_id)
    if not bulletin:
        raise ValueError(f"Boletim id={bulletin_id} nao encontrado.")

    template_path = resolve_template(template_name)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = Path(output_path) if output_path else OUTPUT_DIR / f"boletim_{_safe_filename(bulletin.date_text)}_{bulletin_id}.docx"

    document = Document(str(template_path))
    mapping = build_placeholder_mapping(bulletin)
    replace_placeholders_in_docx(document, mapping)
    leftovers = find_unreplaced_placeholders(document)
    document.save(path)

    if leftovers:
        logger.warning("Boletim DOCX gerado com placeholders nao substituidos: %s", ", ".join(leftovers))
    logger.info("Boletim DOCX gerado com template %s: %s", template_path.name, path)
    return path


def generate_bulletin_by_date(
    date_text: str,
    output_path: str | Path | None = None,
    template_name: str | None = None,
) -> Path:
    database = Database()
    bulletin = database.get_bulletin_by_date(date_text)
    if not bulletin or bulletin.id is None:
        raise ValueError(f"Nenhum boletim encontrado para a data: {date_text}")
    return generate_bulletin_docx(bulletin.id, output_path, template_name=template_name)
