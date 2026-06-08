from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

from src.config import OUTPUT_DIR, ensure_directories
from src.generated_scales_repository import GeneratedScalesRepository, SavedScale


def _safe_filename(value: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9_-]+", "_", value.strip())
    clean = re.sub(r"_+", "_", clean).strip("_")
    return clean or "escala"


def _scale_file_base(scale: SavedScale) -> str:
    return f"escala_{scale.id}_{_safe_filename(scale.title)}"


def export_saved_scale_xlsx(scale_id: int, output_dir: Path | None = None) -> Path:
    ensure_directories()
    repository = GeneratedScalesRepository()
    scale = repository.get_scale(scale_id)
    if not scale:
        raise ValueError("Escala nao encontrada.")
    assignments = repository.get_assignment_rows(scale_id)
    target_dir = output_dir or OUTPUT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / f"{_scale_file_base(scale)}.xlsx"

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Escala"

    sheet["A1"] = scale.title or "Escala"
    sheet["A1"].font = Font(size=16, bold=True)
    sheet.merge_cells("A1:E1")
    sheet["A2"] = "Data/periodo"
    sheet["B2"] = scale.service_date
    sheet["A3"] = "Modelo"
    sheet["B3"] = scale.model_name
    sheet["A4"] = "Status"
    sheet["B4"] = scale.status

    headers = ["Funcao", "Integrante", "Motivo", "Aviso", "Ordem"]
    header_row = 6
    for column_index, header in enumerate(headers, start=1):
        cell = sheet.cell(row=header_row, column=column_index, value=header)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")

    for row_index, assignment in enumerate(assignments, start=header_row + 1):
        sheet.cell(row=row_index, column=1, value=assignment.function_name)
        sheet.cell(row=row_index, column=2, value=assignment.person_name or "A definir")
        sheet.cell(row=row_index, column=3, value=assignment.reason)
        sheet.cell(row=row_index, column=4, value=assignment.warning)
        sheet.cell(row=row_index, column=5, value=assignment.sort_order)

    widths = [24, 28, 42, 42, 10]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    workbook.save(output_path)
    return output_path


def export_saved_scale_docx(scale_id: int, output_dir: Path | None = None) -> Path:
    ensure_directories()
    repository = GeneratedScalesRepository()
    scale = repository.get_scale(scale_id)
    if not scale:
        raise ValueError("Escala nao encontrada.")
    assignments = repository.get_assignment_rows(scale_id)
    target_dir = output_dir or OUTPUT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / f"{_scale_file_base(scale)}.docx"

    document = Document()
    title = document.add_heading(scale.title or "Escala", level=1)
    title.runs[0].font.size = Pt(18)

    if scale.service_date:
        document.add_paragraph(f"Data/periodo: {scale.service_date}")
    if scale.model_name:
        document.add_paragraph(f"Modelo: {scale.model_name}")
    document.add_paragraph(f"Status: {scale.status}")

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Funcao", "Integrante", "Motivo", "Aviso"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for assignment in assignments:
        cells = table.add_row().cells
        cells[0].text = assignment.function_name
        cells[1].text = assignment.person_name or "A definir"
        cells[2].text = assignment.reason
        cells[3].text = assignment.warning

    document.add_paragraph("")
    document.add_paragraph("Revise a escala antes de publicar.")
    document.save(output_path)
    return output_path
