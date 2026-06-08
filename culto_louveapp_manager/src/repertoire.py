from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt

from src.config import OUTPUT_DIR
from src.database import Database
from src.logger import get_logger
from src.models import WorshipBulletin, WorshipSong


logger = get_logger(__name__)


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", value or "sem_data").strip("_")
    return cleaned or "sem_data"


def _songs_from_bulletin(bulletin: WorshipBulletin) -> list[WorshipSong]:
    songs = [
        WorshipSong(bulletin.preludio_musica, bulletin.preludio_cantor, bulletin.preludio_tom, "Preludio"),
        WorshipSong(bulletin.musica1, bulletin.cantor1, bulletin.tom1, "Louvor 1"),
        WorshipSong(bulletin.musica2, bulletin.cantor2, bulletin.tom2, "Louvor 2"),
        WorshipSong(bulletin.musica3, bulletin.cantor3, bulletin.tom3, "Louvor 3"),
        WorshipSong(bulletin.musica4, bulletin.cantor4, bulletin.tom4, "Musica 4"),
        WorshipSong(bulletin.musica5, bulletin.cantor5, bulletin.tom5, "Musica 5"),
        WorshipSong(bulletin.musica_pao, bulletin.cantor_pao, bulletin.tom_pao, "Santa Ceia - Pao"),
        WorshipSong(bulletin.musica_vinho, bulletin.cantor_vinho, bulletin.tom_vinho, "Santa Ceia - Vinho"),
        WorshipSong(bulletin.musica_extra, bulletin.cantor_extra, bulletin.tom_extra, "Extra"),
        WorshipSong(bulletin.musica_final, bulletin.cantor_final, bulletin.tom_final, "Final"),
    ]
    return [song for song in songs if song.title or song.singer or song.key]


def generate_repertoire_docx(bulletin_id: int, output_path: str | Path | None = None) -> Path:
    database = Database()
    bulletin = database.get_bulletin(bulletin_id)
    if not bulletin:
        raise ValueError(f"Boletim id={bulletin_id} nao encontrado.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = Path(output_path) if output_path else OUTPUT_DIR / f"repertorio_{_safe_filename(bulletin.date_text)}_{bulletin_id}.docx"
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    document.add_heading("REPERTORIO DO CULTO", 0)
    document.add_paragraph(f"Data: {bulletin.date_text}")
    document.add_paragraph(f"Dirigente: {bulletin.dirigente}")
    document.add_paragraph(f"Pregador: {bulletin.pregador}")

    document.add_heading("Lista de musicas", 1)
    songs = _songs_from_bulletin(bulletin)
    for song in songs:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(song.section).bold = True
        paragraph.add_run(f": {song.title} | Cantor: {song.singer} | Tom: {song.key}")

    document.add_heading("Tabela", 1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("Secao", "Musica", "Cantor", "Tom")
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for song in songs:
        row = table.add_row().cells
        row[0].text = song.section
        row[1].text = song.title
        row[2].text = song.singer
        row[3].text = song.key

    document.save(path)
    logger.info("Repertorio DOCX gerado: %s", path)
    return path


def generate_repertoire_by_date(date_text: str, output_path: str | Path | None = None) -> Path:
    database = Database()
    bulletin = database.get_bulletin_by_date(date_text)
    if not bulletin or bulletin.id is None:
        raise ValueError(f"Nenhum boletim encontrado para a data: {date_text}")
    return generate_repertoire_docx(bulletin.id, output_path)
